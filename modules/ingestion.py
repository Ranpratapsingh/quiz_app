"""
modules/ingestion.py
--------------------
Handles document ingestion for supported file types and raw pasted text.
Extracts raw text and splits it into chunks suitable for LLM input.

Supported file formats: PDF, TXT, DOCX, CSV
"""

import re
import csv
from pathlib import Path


# ── Chunk settings ─────────────────────────────────────────────────────────────
CHUNK_SIZE   = 800   # characters per chunk
CHUNK_OVERLAP = 100  # overlap between adjacent chunks


# ══════════════════════════════════════════════════════════════════════════════
#  Extractors
# ══════════════════════════════════════════════════════════════════════════════

def extract_pdf(path: str) -> str:
    """Extract text from a PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        doc.close()
        return "\n\n".join(pages)
    except ImportError:
        raise ImportError(
            "PyMuPDF not installed. Run: pip install pymupdf"
        )
    except Exception as e:
        raise RuntimeError(f"PDF extraction failed: {e}")


def extract_txt(path: str) -> str:
    """Extract text from a plain text file."""
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return Path(path).read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    raise RuntimeError("Could not decode text file with common encodings.")


def extract_docx(path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except ImportError:
        raise ImportError(
            "python-docx not installed. Run: pip install python-docx"
        )
    except Exception as e:
        raise RuntimeError(f"DOCX extraction failed: {e}")


def extract_csv(path: str) -> str:
    """
    Extract content from a CSV file.
    Converts to a readable prose-like format so the LLM can reason about rows.
    """
    try:
        rows = []
        with open(path, newline="", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            headers = reader.fieldnames or []
            rows = list(reader)

        if not rows:
            return ""

        # Build a textual summary the LLM can work with
        lines = [f"This document is a CSV table with columns: {', '.join(headers)}.\n"]
        lines.append(f"It has {len(rows)} data rows.\n")

        # First 100 rows as readable text
        for i, row in enumerate(rows[:100]):
            entry = "; ".join(f"{k}: {v}" for k, v in row.items() if v.strip())
            lines.append(f"Row {i+1}: {entry}")

        if len(rows) > 100:
            lines.append(f"... and {len(rows) - 100} more rows.")

        return "\n".join(lines)
    except Exception as e:
        raise RuntimeError(f"CSV extraction failed: {e}")


# ══════════════════════════════════════════════════════════════════════════════
#  Text cleaning
# ══════════════════════════════════════════════════════════════════════════════

def clean_text(text: str) -> str:
    """Remove noise: excessive whitespace, control characters, repeated punctuation."""
    # Remove control characters except newlines/tabs
    text = re.sub(r"[^\S\n\t]+", " ", text)
    # Collapse 3+ newlines into 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove lines that are pure symbols or page numbers
    lines = text.splitlines()
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned.append("")
            continue
        # Skip lines that are only numbers (page numbers)
        if re.fullmatch(r"[\d\s\.\-]+", stripped) and len(stripped) < 6:
            continue
        # Skip lines with only symbols
        if re.fullmatch(r"[^a-zA-Z0-9]+", stripped):
            continue
        cleaned.append(stripped)
    return "\n".join(cleaned).strip()


# ══════════════════════════════════════════════════════════════════════════════
#  Chunking
# ══════════════════════════════════════════════════════════════════════════════

def split_into_chunks(text: str, chunk_size: int = CHUNK_SIZE,
                      overlap: int = CHUNK_OVERLAP) -> list[str]:
    """
    Split text into overlapping chunks.
    Tries to split on paragraph boundaries first; falls back to character splits.
    """
    # Split on paragraphs
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]

    chunks = []
    current = ""

    for para in paragraphs:
        # If adding this paragraph exceeds chunk_size, flush current
        if len(current) + len(para) + 2 > chunk_size and current:
            chunks.append(current.strip())
            # Keep overlap: take last `overlap` chars as the start of next chunk
            current = current[-overlap:] + "\n\n" + para
        else:
            current = (current + "\n\n" + para).strip() if current else para

    if current.strip():
        chunks.append(current.strip())

    # If a single chunk is still too long (e.g. one massive paragraph), split by sentences
    final_chunks = []
    for chunk in chunks:
        if len(chunk) <= chunk_size * 1.5:
            final_chunks.append(chunk)
        else:
            # Split by sentence endings
            sentences = re.split(r"(?<=[.!?])\s+", chunk)
            sub = ""
            for sent in sentences:
                if len(sub) + len(sent) > chunk_size and sub:
                    final_chunks.append(sub.strip())
                    sub = sub[-overlap:] + " " + sent
                else:
                    sub = (sub + " " + sent).strip()
            if sub.strip():
                final_chunks.append(sub.strip())

    return [c for c in final_chunks if len(c) > 50]  # drop tiny fragments


# ══════════════════════════════════════════════════════════════════════════════
#  Public API
# ══════════════════════════════════════════════════════════════════════════════

def ingest_document(path: str) -> list[str]:
    """
    Main entry point.
    Detects file type, extracts text, cleans it, and returns a list of chunks.

    Args:
        path: Path to the uploaded document.

    Returns:
        List of text chunks ready for LLM consumption.

    Raises:
        ValueError: If the file type is unsupported.
        RuntimeError: If extraction fails.
    """
    ext = Path(path).suffix.lower()

    extractors = {
        ".pdf":  extract_pdf,
        ".txt":  extract_txt,
        ".docx": extract_docx,
        ".csv":  extract_csv,
    }

    if ext not in extractors:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            f"Supported: {', '.join(extractors.keys())}"
        )

    raw_text = extractors[ext](path)

    if not raw_text or not raw_text.strip():
        raise RuntimeError(
            "No text could be extracted from the document. "
            "The file may be empty, image-only, or corrupted."
        )

    cleaned = clean_text(raw_text)
    chunks  = split_into_chunks(cleaned)

    if not chunks:
        raise RuntimeError("Text was extracted but chunking produced no usable segments.")

    return chunks


def ingest_text(text: str) -> list[str]:
    """
    Process raw pasted text using the same cleaning and chunking pipeline
    used for uploaded documents.
    """
    if not text or not text.strip():
        raise RuntimeError("No text was provided.")

    cleaned = clean_text(text)
    if not cleaned:
        raise RuntimeError("The pasted text did not contain usable content.")

    chunks = split_into_chunks(cleaned)
    if not chunks:
        raise RuntimeError("The pasted text was too short to create usable chunks.")

    return chunks


# ── Quick self-test ────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python -m modules.ingestion <path_to_file>")
        sys.exit(1)
    result = ingest_document(sys.argv[1])
    print(f"Extracted {len(result)} chunks.")
    print(f"\n--- Chunk 1 preview ---\n{result[0][:400]}")
